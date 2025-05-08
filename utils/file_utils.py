import pdfplumber
from docx import Document as DocxDocument

def extract_text(file_path):
    try:
        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return {'text': text, 'page_count': len(text) // Config.PAGE_SIZE + 1}
        
        elif file_path.endswith('.pdf'):
            with pdfplumber.open(file_path) as pdf:
                text = '\n'.join(page.extract_text() for page in pdf.pages)
                return {'text': text, 'page_count': len(pdf.pages)}
        
        elif file_path.endswith('.docx'):
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            # Подсчет страниц через разрывы (если есть)
            page_count = 1
            for p in doc.paragraphs:
                if 'PAGE' in p.text.upper() and 'BREAK' in p.text.upper():
                    page_count += 1
            return {'text': '\n'.join(paragraphs), 'page_count': page_count}
        
        return None
    except Exception:
        return None